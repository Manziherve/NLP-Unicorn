import os
from dotenv import load_dotenv
from flask import Flask, render_template, jsonify, request, send_file
import json
from io import BytesIO
from docx import Document
from services.comparator_service import ComparatorService
from services.extractor_service import ExtractorService
from services.generate_content import generate_copy

# Initialize Flask application
app = Flask(__name__)

# Load environment variables and initialize services
try: 
    load_dotenv()
    comparator_service = ComparatorService(api_key=os.getenv('GEMINI_API_KEY'))
    extractor_service = ExtractorService(api_key=os.getenv('GEMINI_API_KEY'))
except Exception as e:
    print(f"Error loading environment variables: {e}")
    comparator_service = None
    extractor_service = None

# ========================================================================
# HTML ROUTE HANDLERS
# ========================================================================

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/index')
def index():
    return render_template('index.html')

@app.route('/copyfile')
def copyfile():
    return render_template('copyfile.html')

@app.route('/comparefiles')
def comparefiles():
    return render_template('comparefiles.html')

@app.route('/index.html')
def index_alias():
    return render_template('index.html')

@app.route('/copyfile.html')
def copy_file():
    return render_template('copyfile.html')

@app.route('/comparefiles.html')
def compare_files():
    return render_template('comparefiles.html')

# ========================================================================
# API ENDPOINTS - NEW WORKFLOW
# ========================================================================

# 8888888 Generate copy from briefing text (index.html workflow)
@app.route('/api/generate-copy', methods=['POST'])
def generate_copy_api():
    try:
        data = request.get_json()
        briefing_text = data.get('briefing', '')
        
        if not briefing_text.strip():
            return jsonify({'error': 'No briefing text provided'}), 400
        
        lines = briefing_text.split('\n')
        title = lines[0] if lines else "Marketing Copy"
        
        generated_copy = f"""# {title}

## Executive Summary
Based on your briefing, we've developed compelling marketing copy that captures your brand message and resonates with your target audience.

## Key Benefits
• Engaging and conversion-focused content
• Brand-aligned messaging
• Target audience optimization
• Clear call-to-action integration

## Marketing Copy

{briefing_text[:500]}

This copy is designed to drive engagement and conversions while maintaining brand consistency.

## Call to Action
Ready to take the next step? Contact us today to learn more.

---
Generated from briefing content"""
        
        return jsonify({'copy': generated_copy})
    
    except Exception as e:
        return jsonify({'error': f'Failed to generate copy: {str(e)}'}), 500

# 8888888 Compare briefing and copy content (index.html metrics)
@app.route('/api/compare-content', methods=['POST'])
def compare_content():
    try:
        data = request.get_json()
        briefing = data.get('briefing', '')
        copy_text = data.get('copy', '')
        
        if not briefing or not copy_text:
            return jsonify({'error': 'Both briefing and copy required'}), 400
        
        word_count = len(copy_text.split())
        relevance = min(90, len(set(briefing.lower().split()) & set(copy_text.lower().split())) * 5)
        readability = min(95, 70 + (word_count // 10))
        engagement = min(88, relevance + (word_count // 20))
        
        return jsonify({
            'wordCount': word_count,
            'relevance': relevance,
            'readability': readability,
            'engagement': engagement
        })
    
    except Exception as e:
        return jsonify({'error': f'Comparison failed: {str(e)}'}), 500

# 8888888 Generate design from copy (copyfile.html workflow)
@app.route('/api/generate-design', methods=['POST'])
def generate_design():
    try:
        data = request.get_json()
        copy_text = data.get('copy', '')
        
        if not copy_text.strip():
            return jsonify({'error': 'No copy text provided'}), 400
        
        lines = copy_text.split('\n')
        title = lines[0].replace('#', '').strip() if lines else "Marketing Content"
        
        designed_copy = f"""
    <div style="font-family: 'Roboto', sans-serif; max-width: 100%; background: linear-gradient(135deg, #ffffff 0%, #f8f9fa 100%); border-radius: 16px; overflow: hidden; box-shadow: 0 8px 24px rgba(0, 0, 0, 0.12);">
        <header style="background: linear-gradient(135deg, #ff6c00 0%, #e55a00 100%); color: white; padding: 2rem; text-align: center;">
            <h1 style="font-size: 2rem; font-weight: 700; margin: 0;">{title}</h1>
            <div style="width: 60px; height: 4px; background: rgba(255,255,255,0.8); margin: 1rem auto 0; border-radius: 2px;"></div>
        </header>
        
        <main style="padding: 2rem;">
            <div style="margin-bottom: 2rem;">
                <p style="font-size: 1.1rem; line-height: 1.6; color: #333;">{copy_text[:300]}...</p>
            </div>
            
            <div style="background: rgba(255,108,0,0.05); padding: 1.5rem; border-radius: 12px; margin-bottom: 2rem;">
                <h2 style="color: #ff6c00; margin-bottom: 1rem;">Key Features</h2>
                <ul style="color: #555;">
                    <li>Professional marketing content</li>
                    <li>Brand-aligned messaging</li>
                    <li>Conversion-optimized design</li>
                </ul>
            </div>
            
            <div style="text-align: center;">
                <button style="background: #28a745; color: white; border: none; padding: 1rem 2rem; font-size: 1.1rem; font-weight: 600; border-radius: 8px; cursor: pointer; box-shadow: 0 4px 8px rgba(0,0,0,0.2);">Get Started Today</button>
            </div>
        </main>
        
        <footer style="background: #f8f9fa; padding: 1rem; text-align: center; border-top: 1px solid #e9ecef;">
            <p style="color: #6c757d; margin: 0;">Professional marketing content designed for results</p>
        </footer>
    </div>
    """
        
        return jsonify({'designed_copy': designed_copy})
    
    except Exception as e:
        return jsonify({'error': f'Failed to generate design: {str(e)}'}), 500

# 8888888 Compare files multi-format (comparefiles.html workflow)
@app.route('/api/compare-files', methods=['POST'])
def compare_files_api():
    try:
        data = request.get_json()
        file1 = data.get('file1', {})
        file2 = data.get('file2', {})
        
        if not file1.get('content') or not file2.get('content'):
            return jsonify({'error': 'Both files content required'}), 400
        
        same_type = file1.get('type') == file2.get('type')
        base_score = 70 if same_type else 45
        
        metrics = {
            'similarity': base_score + (len(set(file1['content'].split()) & set(file2['content'].split())) // 10),
            'contentMatch': base_score + (len(file1['content'].split()) // 20),
            'structure': base_score + 15,
            'compatibility': 95 if same_type else 60
        }
        
        return jsonify(metrics)
    
    except Exception as e:
        return jsonify({'error': f'File comparison failed: {str(e)}'}), 500

# 8888888 Generate content from file comparison (comparefiles.html generation)
@app.route('/api/generate-from-comparison', methods=['POST'])
def generate_from_comparison():
    try:
        data = request.get_json()
        file1_content = data.get('file1Content', '')
        file2_content = data.get('file2Content', '')
        generation_type = data.get('generationType', 'copy')
        
        if not file1_content or not file2_content:
            return jsonify({'error': 'Both file contents required'}), 400
        
        combined_content = f"{file1_content}\n\n{file2_content}"
        
        if generation_type == 'design':
            generated_content = f"""
            <div style="font-family: 'Roboto', sans-serif; background: linear-gradient(135deg, #f8f9fa 0%, #fff 100%); border-radius: 12px; overflow: hidden; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
                <header style="background: linear-gradient(135deg, #ff6c00 0%, #e55a00 100%); color: white; padding: 2rem; text-align: center;">
                    <h1>Combined Marketing Content</h1>
                </header>
                
                <main style="padding: 2rem;">
                    <div style="margin-bottom: 2rem; padding: 1rem; background: rgba(255,108,0,0.05); border-radius: 8px;">
                        <h2>Combined Content:</h2>
                        <p>{combined_content[:200]}...</p>
                    </div>
                    
                    <div style="text-align: center;">
                        <button style="background: #28a745; color: white; border: none; padding: 1rem 2rem; border-radius: 8px; font-weight: 600; cursor: pointer;">Learn More</button>
                    </div>
                </main>
            </div>
            """
        else:
            generated_content = f"""Based on the comparison of your files:

{combined_content[:500]}...

This generated copy combines elements from both files to create compelling marketing content."""
        
        return jsonify({'content': generated_content})
    
    except Exception as e:
        return jsonify({'error': f'Generation from comparison failed: {str(e)}'}), 500

# ========================================================================
# API ENDPOINTS - LEGACY (Existing functionality)
# ========================================================================

@app.route('/api/generate_copy', methods=['POST'])
def generate_copy_route():
    if 'doc1' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    doc1 = request.files['doc1']
    
    try:
        decoded_output, _ = generate_copy(doc1)
        return jsonify({'copy': decoded_output})
    
    except Exception as e:
        return jsonify({'error': f'Failed to generate copy: {str(e)}'}), 500

@app.route('/api/download-copy', methods=['POST'])
def download_copy():
    try:
        data = request.get_json()
        copy_text = data.get('copy', '')
        
        if not copy_text:
            return jsonify({'error': 'No copy content provided'}), 400
        
        doc = Document()
        doc.add_heading('Generated Marketing Copy', 0)
        
        paragraphs = copy_text.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return send_file(
            doc_buffer,
            as_attachment=True,
            download_name='generated-copy.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': f'Error creating DOCX: {str(e)}'}), 500

@app.route('/api/download-design', methods=['POST'])
def download_design():
    try:
        data = request.get_json()
        designed_copy = data.get('designed_copy', '')
        
        if not designed_copy:
            return jsonify({'error': 'No design content provided'}), 400
        
        doc = Document()
        doc.add_heading('Generated Marketing Design', 0)
        doc.add_paragraph(designed_copy)
        
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return send_file(
            doc_buffer,
            as_attachment=True,
            download_name='designed-copy.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': f'Error creating design DOCX: {str(e)}'}), 500

@app.route('/api/extract', methods=['POST'])
def extract():
    if 'doc1' not in request.files or 'doc2' not in request.files:
        return jsonify({'error': 'Two files required for extraction'}), 400
    
    doc1 = request.files.getlist('doc1')
    doc2 = request.files.getlist('doc2')
    
    words_to_anonymize = request.form.get('words_to_anonymize', '[]')

    try:
        if isinstance(words_to_anonymize, str):
            words_to_anonymize = json.loads(words_to_anonymize)
    except (json.JSONDecodeError, TypeError):
        words_to_anonymize = []
    
    if not extractor_service:
        return jsonify({'error': 'Extraction service not available'}), 503
        
    result = extractor_service.extract_anonymized(
        doc1, doc2, 
        words_to_anonymize=words_to_anonymize
    )
    
    if result['success']:
        return jsonify({
            'docs': result['docs'],
            'mapping': result['mapping']
        })
    else:
        return jsonify({'error': result['error']}), result.get('status_code', 500)

@app.route('/api/compare', methods=['POST'])
def compare():
    if 'doc1' not in request.files or 'doc2' not in request.files:
        return jsonify({'error': 'Two files required for comparison'}), 400
    
    doc1 = request.files.getlist('doc1')
    doc2 = request.files.getlist('doc2')
    
    words_to_anonymize = request.form.get('words_to_anonymize', '[]')  
    comparison_type = request.form.get('comparison_type', 'copy_design')  
    
    try:
        if isinstance(words_to_anonymize, str):
            words_to_anonymize = json.loads(words_to_anonymize)
    except (json.JSONDecodeError, TypeError):
        words_to_anonymize = []

    if not extractor_service or not comparator_service:
        return jsonify({'error': 'Comparison services not available'}), 503

    result = extractor_service.extract_anonymized(
        doc1, doc2, 
        words_to_anonymize=words_to_anonymize
    )

    if result['success']:
        comp_result = comparator_service.compare(
            result['docs'][0], 
            result['docs'][1], 
            mapping=result['mapping'],
            comparison_type=comparison_type
        )

        if comp_result['success']:
            return jsonify({'result': comp_result['result']})
        else:
            return jsonify({'error': comp_result['error']}), comp_result.get('status_code', 500)
    else:
        return jsonify({'error': result['error']}), result.get('status_code', 500)

# ========================================================================
# APPLICATION CONFIGURATION
# ========================================================================

if __name__ == '__main__':
    app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
    app.run(debug=True, port=5000)