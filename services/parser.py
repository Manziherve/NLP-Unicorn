from docx import Document
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader
import io
from werkzeug.datastructures import FileStorage

from llm.gemini_client import ImageExtractionClient


class FileParser:
	"""
	A class for parsing different file types (DOCX, HTML, images) and cleaning the extracted text.
	"""
	
	def __init__(self, gemini_client: ImageExtractionClient =None):
		"""
		Initialize the FileParser.
				Args:
			gemini_client (GeminiClient, optional): Initialized Gemini client for image parsing
		"""
		self.gemini_client = gemini_client
		

	def parse_docx(self, file_input):
		"""
		Parse a DOCX file and return cleaned text content.
		
		Args:
			file_input (str or FileStorage): Path to the DOCX file or FileStorage object
			
		Returns:
			str: Cleaned text content from the DOCX file
		"""
		try:
			# Handle FileStorage object - read content into memory
			if isinstance(file_input, FileStorage):
				file_input.stream.seek(0)
				doc = Document(file_input.stream)

			else: # local docx file
				doc = Document(file_input)
			
			text_content = []
			
			# Extract text from paragraphs
			for paragraph in doc.paragraphs:
				if paragraph.text.strip():
					text_content.append(paragraph.text)
			
			# Extract text from tables
			for table in doc.tables:
				for row in table.rows:
					row_text = []
					for cell in row.cells:
						if cell.text.strip():
							row_text.append(cell.text)
					if row_text:
						text_content.append(' | '.join(row_text))
			
			raw_text = '\n'.join(text_content)
			return raw_text.strip()
		
		except Exception as e:
			if isinstance(file_input, FileStorage):
				name = file_input.filename
			else:
				name = str(file_input)
			print(f"Error parsing DOCX file {name}: {str(e)}")
			return ""
		
	def parse_pdf(self, file_input):
		"""
		Parse a PDF file and return cleaned text content.
		
		Args:
			file_input (str or FileStorage): Path to the PDF file or FileStorage object
			
		Returns:
			str: Cleaned text content from the PDF file
		"""
		try:
			file_input.seek(0)  # Ensure we're at the beginning
			reader = PdfReader(file_input)
			
			text_content = []
			
			for page in reader.pages:
				if page.extract_text():
					text_content.append(page.extract_text())
			
			raw_text = '\n'.join(text_content)

			return raw_text.strip()
		except Exception as e:
			print(f"Error parsing PDF file {file_input.filename}: {str(e)}")
			return ""
	
	def parse_html(self, file_input):
		"""
		Parse an HTML file and return cleaned text content.
		
		Args:
			file_input (str or FileStorage): Path to the HTML file or FileStorage object
			
		Returns:
			str: Cleaned text content from the HTML file
		"""
		try:
			file_input.seek(0)  # Ensure we're at the beginning
			html_content = file_input.read().decode('utf-8')

			soup = BeautifulSoup(html_content, "html.parser")
			raw_text = soup.get_text(separator=" ", strip=True)
		
			return raw_text.strip()
		except Exception as e:
			print(f"Error parsing HTML file {file_input.filename}: {str(e)}")
			return ""
	
	def parse_image(self, file_input):
		"""
		Extract text from an image using Google Gemini Vision API and return cleaned text.
		
		Args:
			file_input (str or FileStorage): Path to the image file or FileStorage object
			
		Returns:
			str: Cleaned text content extracted from the image
		"""
		try:
			if not self.gemini_client:
				print(f"Gemini client required for image parsing of {file_input.filename}")
				print("Please provide a gemini_client when initializing FileParser")
				return ""
			
			raw_text = self.gemini_client.extract(file_input)
			
			return raw_text.strip()
			
		except Exception as e:
			file_name = file_input.filename if isinstance(file_input, FileStorage) else file_input
			print(f"Error extracting text from image {file_name}: {str(e)}")
			return ""