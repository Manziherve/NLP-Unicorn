import os, asyncio, tempfile, json
from pathlib import Path
from dotenv import load_dotenv
from flask import Flask, render_template, jsonify, request, send_file
import json
from io import BytesIO
from docx import Document
from docx.shared import Inches
from services.comparator_service import ComparatorService
from services.elsa import anonymize_text
from services.extractor_service import ExtractorService
from services.convertor_service import ConvertorService
from services.generate_content import generate_copy, make_filestorage_from
from services.generator_service import GeneratorService

# Initialize Flask application
app = Flask(__name__)

comparator_service = None
generator_service = None  
extractor_service = None
convertor_service = None
# Load environment variables and initialize services
try: 
    load_dotenv()
    comparator_service = ComparatorService(api_key=os.getenv('GEMINI_API_KEY'))
    generator_service = GeneratorService(api_key=os.getenv('GEMINI_API_KEY'))
    extractor_service = ExtractorService(api_key=os.getenv('GEMINI_API_KEY'))
    convertor_service = ConvertorService()
    
except Exception as e:
    print(f"❌ CRITICAL ERROR: {str(e)}")
    import traceback
    traceback.print_exc()

# Routes principales - pages
@app.route('/')
@app.route('/index')
def index():
    return render_template('index.html')

@app.route('/copyfile')
def copyfile():
    return render_template('copyfile.html')

@app.route('/comparefiles')
def comparefiles():
    return render_template('comparefiles.html')

@app.route('/api/generate_copy', methods=['POST'])
def generate_copy_route():
    if 'doc1' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    # words_to_anonymize = request.form.get('words_to_anonymize', '[]')

    doc1 = request.files['doc1']
    decoded_output, _ = generate_copy(doc1) # add more_words
    return jsonify({'output': decoded_output})

@app.route('/api/generate_docx_preview', methods=['POST'])
def generate_docx_preview():
    try:
        if not request.is_json:
            return jsonify({'error': 'Content-Type must be application/json'}), 400
        
        data = request.get_json()
        copy_text = data.get('copy', '')
        
        if not copy_text:
            return jsonify({'error': 'No copy content provided'}), 400
        
        # Créer le document DOCX avec formatage
        doc = Document()
        
        # Configuration des marges
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # Titre principal
        title = doc.add_heading('Generated Marketing Copy', 0)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Espacement
        doc.add_paragraph()
        
        # Contenu avec formatage intelligent
        paragraphs = copy_text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph()
                
                # Détecter les titres (courts et en majuscules)
                if len(para_text.strip()) < 100 and para_text.strip().isupper():
                    # Style titre
                    run = para.add_run(para_text.strip())
                    run.bold = True
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif para_text.strip().endswith(':') and len(para_text.strip()) < 50:
                    # Style sous-titre
                    run = para.add_run(para_text.strip())
                    run.bold = True
                else:
                    # Style paragraphe normal
                    para.add_run(para_text.strip())
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Sauvegarder en mémoire
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=False,
            download_name='preview.docx'
        )
        
    except Exception as e:
        print(f"Error in generate_docx_preview: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/download-copy', methods=['POST'])
def download_copy():
    try:
        if not request.is_json:
            return jsonify({'error': 'Content-Type must be application/json'}), 400
        
        data = request.get_json()
        copy_text = data.get('copy', '')
        
        if not copy_text:
            return jsonify({'error': 'No copy content provided'}), 400
        
        # Réutiliser la même logique que le preview
        doc = Document()
        
        # Configuration des marges
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # Titre principal
        title = doc.add_heading('Generated Marketing Copy', 0)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Espacement
        doc.add_paragraph()
        
        # Contenu avec formatage intelligent
        paragraphs = copy_text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph()
                
                if len(para_text.strip()) < 100 and para_text.strip().isupper():
                    run = para.add_run(para_text.strip())
                    run.bold = True
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif para_text.strip().endswith(':') and len(para_text.strip()) < 50:
                    run = para.add_run(para_text.strip())
                    run.bold = True
                else:
                    para.add_run(para_text.strip())
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Sauvegarder en mémoire pour téléchargement
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return send_file(
            doc_buffer,
            as_attachment=True,
            download_name='generated-copy.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"Error in download_copy: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate_design', methods=['POST'])
def generate_design():
    if 'copy' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    copy = request.files['copy']

    # Get additional parameters from the request body
    words_to_anonymize = request.form.get('words_to_anonymize', '[]')  # Default to empty list if not provided
    generation_type = request.form.get('generation_type', 'design')  # Default to design
    language = request.form.get('language', 'FR') # Default language if not provided
    
    try:
        if isinstance(words_to_anonymize, str):
            words_to_anonymize = json.loads(words_to_anonymize)
    except (json.JSONDecodeError, TypeError):
        # If parsing fails, treat as None
        words_to_anonymize = []

    example_fs = [make_filestorage_from(str(p)) 
                  for p in Path("model_templates/design").glob("*.html")]

    # Extract anonymized content from the documents
    result = extractor_service.extract_anonymized(
        example_fs,
        copy,
        words_to_anonymize=words_to_anonymize,
        parse_html=False
    )

    if result['success']:
        examples = ""
        for idx, example in enumerate(result['docs'][:-1], start=1): 
            examples += f"Example {idx} :\n{example}\n---\n"

        # generate the design using the generator service
        generated_result = generator_service.generate(
            result['docs'][-1], #copy
            mapping=result['mapping'],
            examples=examples,
            generation_type=generation_type,
            language=language
        )

        return generated_result
    else:
        return jsonify({'error': result['error']}), result.get('status_code', 500)


@app.route('/api/extract', methods=['POST'])
def extract():
    if 'doc1' not in request.files or 'doc2' not in request.files:
        return jsonify({'error': 'Two files required for extraction'}), 400
    
    doc1 = request.files.getlist('doc1')
    doc2 = request.files.getlist('doc2')
    
    words_to_anonymize = request.form.get('words_to_anonymize', '[]')

    try:
        if isinstance(words_to_anonymize, str):
            words_to_anonymize = json.loads(words_to_anonymize)
    except (json.JSONDecodeError, TypeError):
        words_to_anonymize = []
    
    if not extractor_service:
        return jsonify({'error': 'Extraction service not available'}), 503
        
    result = extractor_service.extract_anonymized(
        doc1, doc2, 
        words_to_anonymize=words_to_anonymize
    )
    
    if result['success']:
        return jsonify({
            'docs': result['docs'],
            'mapping': result['mapping']
        })
    else:
        return jsonify({'error': result['error']}), result.get('status_code', 500)
  
    
@app.route('/api/compare', methods=['POST'])
def compare():
    try:
        print("🔍 Compare endpoint called")
        print(f"📋 Form keys: {list(request.form.keys())}")
        print(f"📁 File keys: {list(request.files.keys())}")
        
        words_to_anonymize = request.form.get('words_to_anonymize', '[]')  
        comparison_type = request.form.get('comparison_type', 'copy_design')  
        
        try:
            if isinstance(words_to_anonymize, str):
                words_to_anonymize = json.loads(words_to_anonymize)
        except (json.JSONDecodeError, TypeError):
            words_to_anonymize = []

        if not extractor_service or not comparator_service:
            return jsonify({'error': 'Comparison services not available'}), 503
        
        # Get text inputs
        text1 = request.form.get('text1')
        text2 = request.form.get('text2')

        # MODE 1: Si on a du texte, utiliser directement le texte
        if text1 and text2:
            
            # Anonymiser le texte directement
            docs = [
                anonymize_text(text1, words_to_anonymize)[0],
                anonymize_text(text2, words_to_anonymize)[0]
            ]
            mapping = {}  # Pas de mapping pour le texte direct
            
        # MODE 2: Si on a des fichiers, utiliser l'extractor
        elif 'doc1' in request.files and 'doc2' in request.files:
            
            doc1 = [request.files['doc1']]
            doc2 = [request.files['doc2']]
        
            result = extractor_service.extract_anonymized(
                doc1, doc2, 
                words_to_anonymize=words_to_anonymize
            )
            
            if not result['success']:
                print(f"❌ Extractor failed: {result['error']}")
                return jsonify({'error': result['error']}), result.get('status_code', 500)
            
            docs = result['docs']
            mapping = result['mapping']
            
        # MODE 3: Mode mixte - texte ET fichier
        elif (text1 or 'doc1' in request.files) and (text2 or 'doc2' in request.files):
            
            docs = []
            
            # Traiter doc1
            if text1:
                docs.append(anonymize_text(text1, words_to_anonymize)[0])
            else:
                doc1_result = extractor_service.extract_anonymized(
                    [request.files['doc1']], [], 
                    words_to_anonymize=words_to_anonymize
                )
                if not doc1_result['success']:
                    return jsonify({'error': f"Doc1 extraction failed: {doc1_result['error']}"}), 500
                docs.append(doc1_result['docs'][0])
            
            # Traiter doc2
            if text2:
                docs.append(anonymize_text(text2, words_to_anonymize)[0])
            else:
                doc2_result = extractor_service.extract_anonymized(
                    [], [request.files['doc2']], 
                    words_to_anonymize=words_to_anonymize
                )
                if not doc2_result['success']:
                    return jsonify({'error': f"Doc2 extraction failed: {doc2_result['error']}"}), 500
                docs.append(doc2_result['docs'][1])
            
            mapping = {}  # Mapping simplifié pour mode mixte
            
        else:
            return jsonify({'error': 'Please provide either text1/text2 OR doc1/doc2 files'}), 400
        
        comp_result = comparator_service.compare(
            docs[0], 
            docs[1], 
            mapping=mapping,
            comparison_type=comparison_type
        )

        if comp_result['success']:
            return jsonify({'result': comp_result['result']})
        else:
            return jsonify({'error': comp_result['error']}), comp_result.get('status_code', 500)
            
    except Exception as e:
        print(f"❌ Error in compare endpoint: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
    
@app.route('/api/convert', methods=['POST'])
def convert():
    service = ConvertorService()
    # Check if the file is well uploaded
    if 'doc' not in request.files :
        return jsonify({'error': 'No file "doc" provided'}), 400
    doc_zip = request.files['doc']

    # Check if the file is a zip type
    if not service._allowed_file(doc_zip):
        return jsonify({'error' : f'invalid type file {doc_zip.content_type}, must be application/zip'}), 400
    
    # Temporary file to save the ZIP content
    fd_zip, path_zip = tempfile.mkstemp(suffix='.zip')
    os.close(fd_zip) # (we don't need the file descriptor)
    doc_zip.save(path_zip)

    # preparation of the output PDF 
    fd_pdf, path_pdf = tempfile.mkstemp(suffix='.pdf')
    os.close(fd_pdf)

    try:
        asyncio.run(service.html_to_pdf(path_zip, path_pdf))
        return send_file(
            path_pdf,
            as_attachment = True, 
            download_name = 'converted_file.pdf',
            mimetype = 'application/pdf'
        )
    
    except Exception as e:
        return jsonify({'error' : str(e)}), 500
    
    finally:
        for p in (path_zip, path_pdf):
            try: os.remove(p)
            except OSError: pass

if __name__ == '__main__':
    app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
    app.run(debug=True, port=5000)
    # app.run(debug=True, host='0.0.0.0', port=5000)